# import packages
from docx import Document
from docx.shared import Inches
import pyttsx3


# speak fun
def speak(text):
    pyttsx3.speak(text)


# create docx object
doc = Document()

# add picture
doc.add_picture('profile.png', width=Inches(2.0))

doc.add_paragraph("Software Engineer").add_run(' ').bold = True

# take some input data
name = input("Enter Your Name..")
speak('Hi ' + name + ' ho are you today')

speak("Enter Your Phone Number")
phone = input("Enter Your Phone Number..")

speak("Enter Your Email")
email = input("Enter Your Email..")

speak("Enter Your Current Address")
address = input("Enter Your Current Address..")

# append data
doc.add_paragraph(name + ' | ' + phone + ' | ' + email + ' | ' + address)

doc.add_heading("About Me:")
speak("Tell Me about your self..")
doc.add_paragraph(input("Tell Me about your self.."))

doc.add_heading("Education:")
speak("Tell Me about your Education..")
doc.add_paragraph(input("Tell Me about your Education.."))


def experience_section():
    p = doc.add_paragraph()

    speak("Enter Company Name..")

    company = input("Enter Company Name..")
    speak("Enter Start date..")
    start_date = input("Enter Start date..")
    speak("Enter End date..")
    end_date = input("Enter End date..")

    p.add_run(company + ' ').bold = True
    p.add_run(start_date + ' - ' + end_date + '\n').italic = True

    speak("Describe Your Experience at ")
    experience_details = input("Describe Your Experience at " + company)
    p.add_run(experience_details)


doc.add_heading("Work Experience:")

experience_section()

while True:
    speak("Still has experience ? yes/no")
    response = input("Still has experience ? yes/no").lower()
    if response == 'yes':
        experience_section()
    else:
        break

# add skills section
doc.add_heading("Skills:")


def skill_section():
    speak("Enter Your Skill")
    skill = input("Enter Your Skill..")
    p = doc.add_paragraph(skill)
    p.style = 'List Bullet'


skill_section()

while True:
    speak("Still has skills ? yes/no")
    response = input("Still has skills ? yes/no").lower()
    if response == 'yes':
        skill_section()
    else:
        break

# footer
sec = doc.sections[0]
foot = sec.footer
p = foot.paragraphs[0]
p.text = "CV Generated By @Muhammad Salah 27/3/2021"

# save docx
doc.save("cv.docx")
